from io import BytesIO
from PyPDF2 import PdfReader
from fpdf import FPDF
from docx import Document
import re

def save_as_docx(content, file_name):
    """
    Saves content as a DOCX file with grading results.
    """
    doc = Document()
    # Split the content into lines for adding to the doc
    lines = content.split("\n")
    for line in lines:
        doc.add_paragraph(line)
    # Save the document
    doc.save(file_name)


def extract_model_data(file_content):
    """
    Extracts questions and answers from the model Q&A file content.
    """
    model_data = {}
    question_pattern = r"(\d+)\.\s*(.*?)\nAnswer:(.*?)(?=\n\d+\.|$)"
    matches = re.findall(question_pattern, file_content, re.DOTALL)

    for question_number, question_text, answer_text in matches:
        model_data[question_number.strip()] = answer_text.strip()
    return model_data


def extract_student_data(text: str):
    student_data = {}
    questions_and_answers = {}

    # Define multiple patterns to try
    patterns = [
        r"(\d+)\.\s*(.*?)\nAnswer:(.*?)(?=\n\d+\.|$)",  # Default pattern
        r"Question (\d+):\s*(.*?)\s*Answer:\s*(.*?)(?=\s*Question \d+|$)"  # Alternative pattern
    ]

    # Try each pattern until matches are found
    for pattern in patterns:
        matches = re.findall(pattern, text, re.DOTALL)
        if matches:
            # Populate the dictionary with questions and answers
            for question_number, question_text, answer_text in matches:
                questions_and_answers[question_number.strip()] = {
                    "question": question_text.strip(),
                    "answer": answer_text.strip()
                }
            break  # Exit the loop once matches are found

    print(f"questions_and_answers : {questions_and_answers}")
    student_data['questionAndAnswers'] = questions_and_answers
    return student_data


def append_grading_results(student_content, grading_results):
    """
    Appends grading results to the student content after each question and answer.
    """
    updated_text = ""
    question_pattern = r"Question (\d+):\s*(.*?)\nAnswer:\s*(.*?)(?=\nQuestion \d+:|$)"
    question_pattern = r"(\d+)\.\s*(.*?)\nAnswer:(.*?)(?=\n\d+\.|$)"
    last_pos = 0
    total_score = 0
    feedbacks = []
       # Define multiple patterns to try
  
    patterns = [
        r"(\d+)\.\s*(.*?)\nAnswer:(.*?)(?=\n\d+\.|$)",  # Default pattern
        r"Question (\d+):\s*(.*?)\s*Answer:\s*(.*?)(?=\s*Question \d+|$)"  # Alternative pattern
    ]

    # Try each pattern until matches are found
    for pattern in patterns:
        matches = re.findall(pattern, student_content, re.DOTALL)
        if matches:
            last_pos = 0
            for match in matches:
                question_number, question_text, answer_text = match
                start_pos = student_content.find(f"{question_number}. {question_text}\nAnswer:{answer_text}", last_pos)
                if start_pos == -1:
                    start_pos = student_content.find(f"Question {question_number}: {question_text}\nAnswer: {answer_text}", last_pos)
                end_pos = start_pos + len(f"{question_number}. {question_text}\nAnswer:{answer_text}")
                updated_text += student_content[last_pos:start_pos]

                # Append the question, answer, and grading results
                updated_text += f"{question_number}. {question_text}\n"
                updated_text += f"Answer: {answer_text}\n"
                
                # Grading Results for the current question
                grading_result = grading_results.get(question_number.strip(), {})
                total_score += grading_result.get('score_achieved', 0)
                updated_text += f"Score: {grading_result.get('score_achieved', 0)}/{grading_result.get('maximum_score', 0)}\n"
                updated_text += f"Justification: {grading_result.get('justification', 'N/A')}\n"
                updated_text += f"Feedback: {grading_result.get('feedback', 'No feedback provided')}\n"
                feedbacks.append(grading_result.get('feedback', ''))

                last_pos = end_pos
            # Add any remaining content after the last question
            updated_text += student_content[last_pos:]
            break  # Exit the loop once matches are found

    return updated_text, total_score, feedbacks
    

def extract_pdf_text(pdf_content):
    """
    Extracts text from a PDF file.
    """
    reader = PdfReader(BytesIO(pdf_content))
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text


def extract_docx_text(docx_content):
    """
    Extracts text from a DOCX file, including generic numbers from numbered lists and excluding bullet points.
    """
    try:
        # Wrap binary content in a BytesIO object to mimic a file-like object
        doc = Document(BytesIO(docx_content))
        text = []
        numbering_dict = {}

        # Iterate through each paragraph in the document
        for para in doc.paragraphs:
            # Extract the numbering (if present) using the XML
            p_xml = para._element
            numbering = get_generic_numbering_from_xml(p_xml)
            
            # Check if the paragraph has numbering (e.g., in a numbered list)
            if numbering:
                # Check if the paragraph has numbering, and add the number
                if numbering not in numbering_dict:
                    numbering_dict[numbering] = 1
                else:
                    numbering_dict[numbering] += 1
                
                # Add the numbering and the paragraph text
                text.append(f"{numbering_dict[numbering]}. {para.text}")
            else:
                # If no numbering, just add the plain text
                text.append(para.text)

        # Return the cleaned-up text, joining all paragraphs with newlines
        return "\n".join(text).strip()
    
    except Exception as e:
        print(f"Error reading DOCX content: {e}")
        return ""

def get_generic_numbering_from_xml(paragraph_xml):
    """
    Extracts generic numbering information from a paragraph's XML element if available.
    """
    try:
        # Define namespaces for XML parsing
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }
        
        # Get the numbering properties from the paragraph XML element
        num_pr = paragraph_xml.find('.//w:numPr', namespaces)
        
        if num_pr is not None:
            num_id = num_pr.find('.//w:numId', namespaces)
            if num_id is not None:
                num_val = num_id.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                return num_val  # This returns the number ID used in the document (this is the reference ID)
        return None
    except Exception as e:
        print(f"Error extracting numbering: {e}")
        return None

def save_as_text(content, file_name):
    """
    Saves content as a plain text file.
    """
    with open(file_name, "w", encoding="utf-8") as f:
        f.write(content)


def save_as_pdf(content, file_name):
    """
    Saves content as a PDF file with retained font and color.
    """
    pdf = FPDF()
    pdf.add_page()

    print("Saving info as PDF...")

    content = content.encode('latin-1', 'replace').decode('latin-1')

    # Set font (change to your desired font and size)
    pdf.set_font("Arial", size=12)  # You can adjust this

    # Set text color (e.g., RGB for red, blue, etc.)
    pdf.set_text_color(0, 0, 0)  # Black color

    # Split the content into lines for multi-cell format
    lines = content.split("\n")
    
    # Add each line to the PDF, retaining formatting
    for line in lines:
        pdf.multi_cell(0, 10, line)
    
    # Output PDF to file
    pdf.output(file_name)
    
def generate_summary(total_score, max_possible_score, overall_feedback):
    percentage = (total_score / max_possible_score) * 100
    grade = assign_grade(percentage)
    return (
        f"Summary:\n"
        f"Total Score: {total_score}\n"
        f"Percentage: {percentage:.2f}%\n"
        f"Grade: {grade}\n"
        f"Overall Feedback: {overall_feedback}"
    )


def assign_grade(percentage):
    if percentage >= 93:
        return 'A'
    elif percentage >= 90:
        return 'A-'
    elif percentage >= 87:
        return 'B+'
    elif percentage >= 83:
        return 'B'
    elif percentage >= 80:
        return 'B-'
    elif percentage >= 77:
        return 'C+'
    elif percentage >= 73:
        return 'C'
    elif percentage >= 70:
        return 'C-'
    elif percentage >= 67:
        return 'D+'
    elif percentage >= 63:
        return 'D'
    elif percentage >= 60:
        return 'D-'
    else:
        return 'F'
